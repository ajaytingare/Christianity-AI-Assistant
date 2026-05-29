import os
import logging
from typing import List, Optional, Tuple
from pathlib import Path
from app.models.schemas import DocumentChunk, Document
from app.utils.helpers import TextProcessing, FileUtils, IDGenerator
from app.config import get_settings
import asyncio
import time

logger = logging.getLogger(__name__)
settings = get_settings()


class DocumentProcessor:
    """Process uploaded documents (PDF, TXT, DOCX)."""
    
    def __init__(self):
        self.settings = get_settings()
        self.text_processor = TextProcessing()
    
    async def extract_text_from_pdf(self, file_path: str) -> Tuple[str, dict]:
        """Extract text from PDF using PyPDF2."""
        try:
            import PyPDF2
            text = ""
            page_count = 0
            
            with open(file_path, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                page_count = len(pdf_reader.pages)
                
                for page in pdf_reader.pages:
                    text += page.extract_text()
            
            return text, {"pages": page_count, "file_size": os.path.getsize(file_path)}
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise
    
    async def extract_text_from_txt(self, file_path: str) -> Tuple[str, dict]:
        """Extract text from TXT file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            
            return text, {"file_size": os.path.getsize(file_path), "encoding": "utf-8"}
        except Exception as e:
            logger.error(f"Error extracting TXT: {e}")
            raise
    
    async def extract_text_from_docx(self, file_path: str) -> Tuple[str, dict]:
        """Extract text from DOCX file."""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            return text, {"paragraphs": len(doc.paragraphs), "file_size": os.path.getsize(file_path)}
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise
    
    async def extract_text(self, file_path: str, file_type: str) -> Tuple[str, dict]:
        """Extract text from document based on file type."""
        if file_type == "pdf":
            return await self.extract_text_from_pdf(file_path)
        elif file_type == "txt":
            return await self.extract_text_from_txt(file_path)
        elif file_type == "docx":
            return await self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def chunk_text(
        self,
        text: str,
        file_id: str,
        filename: str,
        chunk_size: Optional[int] = None,
        overlap: Optional[int] = None
    ) -> List[DocumentChunk]:
        """Split text into chunks with metadata."""
        chunk_size = chunk_size or self.settings.CHUNK_SIZE
        overlap = overlap or self.settings.CHUNK_OVERLAP
        
        if not text.strip():
            logger.warning(f"Empty text for file {filename}")
            return []
        
        text = self.text_processor.clean_text(text)
        raw_chunks = self.text_processor.generate_chunks(text, chunk_size, overlap)
        
        chunks = []
        for idx, chunk_text in enumerate(raw_chunks):
            chunk = DocumentChunk(
                id=IDGenerator.generate_chunk_id(),
                content=chunk_text,
                source_file=filename,
                page_number=None,
                metadata={
                    "file_id": file_id,
                    "chunk_index": idx,
                    "chunk_size": len(chunk_text),
                    "created_at": str(time.time())
                },
                chunk_index=idx
            )
            chunks.append(chunk)
        
        logger.info(f"Created {len(chunks)} chunks from {filename}")
        return chunks
    
    async def validate_document(self, file_path: str) -> Tuple[bool, Optional[str]]:
        """Validate document before processing."""
        if not os.path.exists(file_path):
            return False, "File does not exist"
        
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_size_mb > self.settings.MAX_FILE_SIZE_MB:
            return False, f"File too large ({file_size_mb:.2f}MB > {self.settings.MAX_FILE_SIZE_MB}MB)"
        
        return True, None
